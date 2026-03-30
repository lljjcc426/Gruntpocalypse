# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

def md_to_docx(md_path, docx_path):
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(11)
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    in_code_block = False
    in_table = False
    table_data = []
    
    i = 0
    while i < len(lines):
        line = lines[i].rstrip('\n')
        
        # 代码块
        if line.startswith('```'):
            in_code_block = not in_code_block
            if in_code_block:
                # 开始代码块
                pass
            i += 1
            continue
        
        if in_code_block:
            p = doc.add_paragraph(line)
            p.style = 'Normal'
            p.paragraph_format.left_indent = Inches(0.3)
            for run in p.runs:
                run.font.name = 'Consolas'
                run.font.size = Pt(10)
            i += 1
            continue
        
        # 表格
        if line.startswith('|'):
            if not in_table:
                in_table = True
                table_data = []
            # 跳过分隔行
            if '---' in line:
                i += 1
                continue
            row = [cell.strip() for cell in line.split('|')[1:-1]]
            table_data.append(row)
            i += 1
            continue
        elif in_table:
            # 表格结束，创建表格
            if table_data:
                num_cols = len(table_data[0])
                table = doc.add_table(rows=len(table_data), cols=num_cols)
                table.style = 'Table Grid'
                for r_idx, row in enumerate(table_data):
                    for c_idx, cell in enumerate(row):
                        if c_idx < num_cols:
                            table.rows[r_idx].cells[c_idx].text = cell
                doc.add_paragraph()
            in_table = False
            table_data = []
        
        # 标题
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
        # 列表
        elif line.startswith('- '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
        # 空行
        elif line.strip() == '':
            pass
        # 普通段落
        else:
            # 处理粗体
            text = line
            p = doc.add_paragraph()
            # 简单处理 **text** 粗体
            parts = re.split(r'(\*\*[^*]+\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
        
        i += 1
    
    # 处理最后的表格
    if in_table and table_data:
        num_cols = len(table_data[0])
        table = doc.add_table(rows=len(table_data), cols=num_cols)
        table.style = 'Table Grid'
        for r_idx, row in enumerate(table_data):
            for c_idx, cell in enumerate(row):
                if c_idx < num_cols:
                    table.rows[r_idx].cells[c_idx].text = cell
    
    doc.save(docx_path)
    print(f'已保存: {docx_path}')

if __name__ == '__main__':
    md_to_docx(
        r'C:\Users\code\Desktop\newweilai\gui\Grunt-2.4.5.250307\Grunt-2.4.5.250307\Web使用手册.md',
        r'C:\Users\code\Desktop\newweilai\gui\Grunt-2.4.5.250307\Grunt-2.4.5.250307\Web使用手册.docx'
    )

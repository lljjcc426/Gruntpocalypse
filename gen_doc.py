#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""生成 Grunt Web UI 使用说明文档"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

doc = Document()

# ---------- 样式调整 ----------
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)
style.paragraph_format.line_spacing = 1.35
style.paragraph_format.space_after = Pt(4)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = '微软雅黑'
    hs.font.color.rgb = RGBColor(0x2B, 0x2B, 0x2B)

# ---------- 封面 ----------
for _ in range(6):
    doc.add_paragraph('')

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Gruntpocalypse Web UI')
run.bold = True
run.font.size = Pt(28)
run.font.name = '微软雅黑'

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run('使用说明  v2.4.5')
r2.font.size = Pt(14)
r2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph('')
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = date_p.add_run('2026.02')
r3.font.size = Pt(11)
r3.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

# ---------- 正文 ----------

doc.add_heading('这是什么', level=1)
doc.add_paragraph(
    'Grunt 本身是一个 JVM 字节码混淆器，之前只有 Swing 桌面版界面。'
    '现在加了一套基于浏览器的 Web 操作界面，功能和桌面版一样，'
    '但用起来更直观——打开浏览器就能用，不用折腾 Swing 那套东西了。'
)
doc.add_paragraph(
    '底层用的 Ktor 起了一个本地 HTTP 服务，默认端口 8080。'
    '前端是纯 HTML/CSS/JS，没有引入任何框架，轻量够用。'
)

doc.add_heading('怎么启动', level=1)

doc.add_heading('最简单的方式', level=2)
doc.add_paragraph(
    '双击 jar 或者命令行直接跑：'
)
p = doc.add_paragraph()
run = p.add_run('java -jar grunt-main.jar')
run.font.name = 'Consolas'
run.font.size = Pt(10)
doc.add_paragraph(
    '不带任何参数的时候，默认就是 Web 模式。'
    '启动后会自动打开浏览器跳转到 http://localhost:8080。'
    '如果浏览器没弹出来，手动访问这个地址就行。'
)

doc.add_heading('指定端口', level=2)
p = doc.add_paragraph()
run = p.add_run('java -jar grunt-main.jar --port=9090')
run.font.name = 'Consolas'
run.font.size = Pt(10)
doc.add_paragraph('端口被占用的时候换一个就好。')

doc.add_heading('还想用老版 Swing 界面？', level=2)
p = doc.add_paragraph()
run = p.add_run('java -jar grunt-main.jar --gui')
run.font.name = 'Consolas'
run.font.size = Pt(10)
doc.add_paragraph('加 --gui 参数走原来的桌面窗口，和以前完全一样。')

doc.add_heading('命令行模式（无界面）', level=2)
p = doc.add_paragraph()
run = p.add_run('java -jar grunt-main.jar config.json')
run.font.name = 'Consolas'
run.font.size = Pt(10)
doc.add_paragraph('直接传一个 .json 配置文件就走 CLI 模式，不开任何界面。适合 CI/CD 或者批量处理。')

doc.add_heading('界面说明', level=1)
doc.add_paragraph(
    '打开网页后整个界面分成几块，类似 IDE 的布局：'
)

doc.add_heading('顶部工具栏', level=2)
doc.add_paragraph(
    '三个主要按钮：「Upload JAR」上传要混淆的 jar 文件，'
    '「Obfuscate」开始混淆，「Download」下载混淆后的结果。'
    '右边有个状态指示灯，显示当前是空闲、就绪、运行中还是已完成。'
)

doc.add_heading('左侧面板 —— 文件树', level=2)
doc.add_paragraph(
    '上传 jar 之后这里会列出里面所有的 class 文件，按包名分成目录结构。'
    '点文件夹可以展开/折叠，点某个 class 就能在中间面板看到反编译后的源码。'
)
doc.add_paragraph(
    '面板顶部有个下拉框可以选择「步骤」——原始代码、某个 Transformer 处理完之后的代码、'
    '以及最终代码。切换步骤后文件树会刷新成对应阶段的类列表。'
)

doc.add_heading('中间面板 —— 代码预览', level=2)
doc.add_paragraph(
    '选中一个 class 之后，这里会展示用 CFR 反编译出来的 Java 源码，带语法高亮。'
    '上面的标签页对应每个处理步骤，可以来回切换对比混淆前后的代码变化。'
)
doc.add_paragraph(
    '说白了就是让你看到每一步 Transformer 到底把代码改成了什么样，'
    '方便调试和调整配置。'
)

doc.add_heading('右侧面板 —— 配置', level=2)
doc.add_paragraph('分两块：')
doc.add_paragraph(
    '上面是「Settings」，可以设置输出文件名、是否开多线程、是否导出映射表、'
    '是否破坏输出格式等。'
)
doc.add_paragraph(
    '下面是「Transformers」，所有可用的混淆器按类别分组排列。'
    '每个混淆器右边有个开关，点一下开启/关闭。'
    '点混淆器的名字区域可以展开它的详细参数进行调整。'
)
doc.add_paragraph(
    '右上角三个小按钮分别是：保存配置到文件、从文件加载配置、重置为默认值。'
)

doc.add_heading('底部面板 —— 控制台', level=2)
doc.add_paragraph(
    '实时显示运行日志，包括读取 jar、各个 Transformer 的执行信息、耗时统计等。'
    '和命令行里看到的输出是一样的内容。右上角有个清空按钮。'
)

doc.add_heading('面板大小调整', level=2)
doc.add_paragraph(
    '各面板之间的分隔线可以拖动，调节左中右三栏和底部控制台的大小。'
    '鼠标放上去会变成蓝色高亮提示。'
)


doc.add_heading('操作流程', level=1)

doc.add_paragraph('正常使用的话，大概就这么几步：')

steps = [
    '点「Upload JAR」选择要混淆的 jar 文件，等上传完成。',
    '右侧面板里勾选要用的 Transformer，按需调整参数。输出文件名默认是 output.jar，可以改。',
    '点「Obfuscate」开始。顶部会出现进度条，控制台实时滚动日志。',
    '混淆完成后，中间面板上方会出现各步骤的标签页，可以点进去看每一步的代码变化。',
    '确认没问题后点「Download」把结果下载下来。',
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {step}')

doc.add_paragraph(
    '如果想换一个 jar 重新来，直接再点 Upload 就行，会自动重置当前会话。'
)


doc.add_heading('Transformer 分类说明', level=1)
doc.add_paragraph('界面上按颜色标签区分了几个类别，简单说一下：')

categories = [
    ('Optimization', '优化类。移除调试信息、精简无用代码、Kotlin 元数据清理等。一般建议开着。'),
    ('Miscellaneous', '杂项。添加垃圾类、克隆类、字段声明混淆、HWID 验证器等。按需选用。'),
    ('Controlflow', '控制流混淆。把 if/switch/循环这些结构打乱，让反编译出来的逻辑变得难读。'),
    ('Encryption', '加密类。字符串加密、数字加密、算术表达式混淆、常量池加密。核心功能。'),
    ('Redirect', '重定向。把字段访问和方法调用替换成间接调用方式（InvokeDynamic 等）。'),
    ('Renaming', '重命名。类名、方法名、字段名、局部变量名全部替换成无意义字符。'),
    ('Minecraft', 'Minecraft 专用。处理 Mixin 注解相关的类名/字段名重命名，MC mod 开发者会用到。'),
]
for name, desc in categories:
    p = doc.add_paragraph()
    run = p.add_run(f'{name}：')
    run.bold = True
    p.add_run(desc)


doc.add_heading('一些注意事项', level=1)

notes = [
    '代码预览依赖 CFR 反编译器，遇到特别复杂或者已经被深度混淆过的字节码时，反编译可能会比较慢或者失败。这不影响混淆本身，只是预览看不了。',
    '语法高亮用的是 CDN 上的 highlight.js，所以预览代码需要能联网。如果在纯内网环境，高亮功能会失效，但代码文本还是能正常显示的。',
    '每个混淆步骤都会在内存里保存一份 class 快照，jar 越大、开的 Transformer 越多，内存占用越高。几十 MB 的 jar 一般没问题，特别大的话注意给 JVM 多分配点内存（-Xmx）。',
    '混淆过程中不要关闭浏览器页面或者终端窗口。浏览器关了的话混淆还是会继续跑完（后台线程），但你看不到进度了，得等它自己结束后重新打开页面。',
    '输出文件默认写到程序运行目录下。如果用的是相对路径（比如 output.jar），就在你执行 java -jar 的那个目录里找。',
    '配置保存/加载功能操作的是服务器端的文件（默认 config.json），不是浏览器本地的。多人用同一个实例的话配置会互相覆盖。正常情况下应该是本机单人使用，不会有这个问题。',
]
for note in notes:
    doc.add_paragraph(note, style='List Bullet')


doc.add_heading('技术细节（给开发者看的）', level=1)

doc.add_paragraph(
    '如果你需要了解底层实现或者想改代码，下面这些信息可能有用：'
)

doc.add_heading('项目结构', level=2)
details = [
    'Web 相关的后端代码在 grunt-main/src/main/kotlin/net/spartanb312/grunt/web/ 下面，一共三个文件：',
]
for d in details:
    doc.add_paragraph(d)

files_desc = [
    ('WebServer.kt', '用 Ktor 搭的 HTTP 服务。负责静态文件托管、REST API（十几个接口）、以及两个 WebSocket 通道（控制台日志和进度推送）。'),
    ('ObfuscationSession.kt', '管理一次混淆会话。核心是在每个 Transformer 执行完之后把所有 ClassNode 序列化成 byte[] 存起来，这样后面才能按步骤回看。'),
    ('Decompiler.kt', 'CFR 反编译器的封装。实现了 ClassFileSource 接口，把内存中的 class 字节喂给 CFR，拿到反编译后的源码字符串。'),
]
for fname, desc in files_desc:
    p = doc.add_paragraph()
    run = p.add_run(f'{fname}')
    run.bold = True
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    p.add_run(f' —— {desc}')

doc.add_paragraph('')
doc.add_paragraph(
    '前端文件在 grunt-main/src/main/resources/web/ 下面：'
    'index.html 是主页面，css/style.css 是样式（Catppuccin 深色主题），'
    'js/api.js 是 API 客户端封装，js/app.js 是界面逻辑。'
)

doc.add_heading('API 一览', level=2)
apis = [
    'GET  /api/transformers       获取所有混淆器及其参数',
    'GET  /api/config              获取当前配置',
    'POST /api/config              更新配置',
    'POST /api/config/save         保存配置到文件',
    'POST /api/config/load         从文件加载配置',
    'POST /api/config/reset        重置配置',
    'POST /api/upload              上传 JAR 文件',
    'POST /api/obfuscate           开始混淆',
    'GET  /api/status              查询混淆状态',
    'GET  /api/steps               获取已完成的步骤列表',
    'GET  /api/classes/{step}      获取某步骤的类列表',
    'GET  /api/preview/{step}/{cls} 获取某步骤某类的反编译源码',
    'GET  /api/download            下载混淆后的 JAR',
    'GET  /api/logs                获取控制台日志',
    'WS   /ws/console              控制台日志实时推送',
    'WS   /ws/progress             混淆进度实时推送',
]
for api in apis:
    p = doc.add_paragraph()
    run = p.add_run(api)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

doc.add_heading('依赖', level=2)
doc.add_paragraph(
    '新增的依赖主要是 Ktor 2.3.7 全家桶（core、netty、content-negotiation、gson、cors、websockets、html-builder）'
    '和 CFR 0.152（Java 反编译器）。都通过 Gradle 的 library 配置打进 fat jar 里了。'
)

doc.add_heading('入口逻辑', level=2)
doc.add_paragraph(
    'Grunt.kt 的 main 函数里加了判断：无参数启动走 Web 模式，'
    '--gui 走原来的 Swing，传 .json 文件走 CLI。'
    '三种模式互不影响，原有功能没有改动。'
)

# ---------- 保存 ----------
output_path = r'C:\Users\code\Desktop\newweilai\gui\Grunt-2.4.5.250307\Grunt-2.4.5.250307\Grunt Web UI 使用说明.docx'
doc.save(output_path)
print(f'文档已生成: {output_path}')
